import os
import io
import re
import base64
from typing import Optional, List, Dict, Any, Tuple
try:
    from django.conf import settings
except ImportError:
    settings = None
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from bs4 import BeautifulSoup
try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None


def natural_sort_key(s: str):
    """Ordenação natural para referências como 4.1, 4.1.1, 4.10, etc."""
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', str(s or ""))]


def set_cell_background(cell, color_hex: str):
    """Aplica cor de fundo a uma célula de tabela docx via XML."""
    color_hex = color_hex.replace("#", "").upper()
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """Define margens internas (padding) de uma célula docx em dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    """Aplica bordas suaves personalizadas a uma tabela docx."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))


def get_template_docx_path() -> str:
    """Retorna o caminho do template docx no servidor."""
    base_dir = getattr(settings, 'BASE_DIR', None)
    if not base_dir:
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
    if os.environ.get('VERCEL') == '1' or str(base_dir).startswith('/var/task'):
        templates_dir = os.path.join('/tmp', 'templates_docx')
    else:
        templates_dir = os.path.join(str(base_dir), 'auditoria', 'templates_docx')
    
    try:
        os.makedirs(templates_dir, exist_ok=True)
    except OSError:
        templates_dir = os.path.join('/tmp', 'templates_docx')
        try:
            os.makedirs(templates_dir, exist_ok=True)
        except OSError:
            pass
    return os.path.join(templates_dir, 'relatorio_auditoria_template.docx')


def ensure_master_template_docx_exists() -> str:
    """
    Cria ou atualiza o arquivo físico de template .docx no servidor com o padrão
    EssilorLuxottica (GQS-POL-017), fontes, tabelas com tags e seções corporativas.
    """
    template_path = get_template_docx_path()
    if os.path.exists(template_path) and os.path.getsize(template_path) > 1000:
        return template_path

    doc = Document()
    for s in doc.sections:
        s.page_height = Inches(11.69)
        s.page_width = Inches(8.27)
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # 1. Header Banner
    h_table = doc.add_table(rows=1, cols=1)
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_h = h_table.cell(0, 0)
    set_cell_background(c_h, "0B2545")
    set_cell_margins(c_h, top=140, bottom=140, left=180, right=180)
    
    p_top = c_h.paragraphs[0]
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_top = p_top.add_run("RELATÓRIO DE AUDITORIA INTERNA DA QUALIDADE")
    r_top.bold = True
    r_top.font.size = Pt(14)
    r_top.font.color.rgb = RGBColor(255, 255, 255)

    p_sub = c_h.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Diretriz Corporativa GQS-POL-017 | {{norma_codigo}} — {{norma_descricao}}")
    r_sub.font.size = Pt(9.5)
    r_sub.font.color.rgb = RGBColor(226, 232, 240)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, color="CBD5E1")
    meta_rows = [
        ("Unidade / Empresa:", "{{unidade}}", "Data da Auditoria:", "{{data_auditoria}}"),
        ("Auditor(es) Líder(es):", "{{auditores}}", "Representantes Auditados:", "{{representantes}}"),
        ("Escopo da Auditoria:", "{{escopo}}", "Modalidade:", "{{modalidade}}"),
        ("Status do Ciclo:", "{{status}}", "Data do Relatório:", "{{data_relatorio}}"),
    ]
    for r_idx, (k1, v1, k2, v2) in enumerate(meta_rows):
        c0 = meta_table.cell(r_idx, 0)
        c1 = meta_table.cell(r_idx, 1)
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        set_cell_background(c0, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
        set_cell_background(c1, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")

        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p0.add_run(f"{k1} ").bold = True
        p0.add_run(v1)

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.add_run(f"{k2} ").bold = True
        p1.add_run(v2)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. Section 1: Avaliação Geral
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run("1. AVALIAÇÃO GERAL E RESULTADOS DA AUDITORIA")
    r1.bold = True
    r1.font.size = Pt(11.5)
    r1.font.color.rgb = RGBColor(11, 37, 69)

    res_table = doc.add_table(rows=8, cols=4)
    res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(res_table, color="CBD5E1")
    
    table_data = [
        ("Classificação", "Sigla", "Qtd", "% Base"),
        ("Conforme", "C", "{{total_c}}", "{{pct_c}}%"),
        ("Não Conformidade Menor", "NC Menor", "{{total_nc_menor}}", "{{pct_nc_menor}}%"),
        ("Não Conformidade Maior", "NC Maior", "{{total_nc_maior}}", "{{pct_nc_maior}}%"),
        ("Oportunidade de Melhoria", "OM", "{{total_om}}", "{{pct_om}}%"),
        ("Observação / Correção Imediata", "OBS", "{{total_obs}}", "{{pct_obs}}%"),
        ("Não Aplicável (fora do escopo)", "NA", "{{total_na}}", "-"),
        ("Total de Requisitos / Amostras Avaliadas", "TOTAL", "{{total_avaliados}}", "100.0%"),
    ]
    for r_idx, (l0, l1, l2, l3) in enumerate(table_data):
        for c_idx, val in enumerate([l0, l1, l2, l3]):
            cell = res_table.cell(r_idx, c_idx)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if r_idx == 0:
                set_cell_background(cell, "0B2545")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx in [1, 2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                r.font.size = Pt(9)
                if r_idx == 7 or c_idx == 1:
                    r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Card Índice & Veredito
    card_t = doc.add_table(rows=1, cols=2)
    card_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(card_t, color="16A34A")
    c_s = card_t.cell(0, 0)
    set_cell_background(c_s, "DCFCE7")
    set_cell_margins(c_s, top=80, bottom=80, left=120, right=120)
    p_cs = c_s.paragraphs[0]
    p_cs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cs.add_run("ÍNDICE GERAL DE CONFORMIDADE\n").bold = True
    p_cs.add_run("{{pct_conformidade}}%\n").bold = True

    c_v = card_t.cell(0, 1)
    set_cell_background(c_v, "F8FAFC")
    set_cell_margins(c_v, top=80, bottom=80, left=120, right=120)
    p_cv = c_v.paragraphs[0]
    p_cv.add_run("PARECER / VEREDITO DA AUDITORIA:\n").bold = True
    p_cv.add_run("{{veredito_status}}\n").bold = True
    p_cv.add_run("{{veredito_parecer}}")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. Section 2: Exclusões Justificadas
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("2. EXCLUSÕES DE REQUISITOS NORMATIVOS (JUSTIFICADAS)")
    r2.bold = True
    r2.font.size = Pt(11.5)
    r2.font.color.rgb = RGBColor(11, 37, 69)

    doc.add_paragraph("{{exclusoes_justificadas}}")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5. Section 3: Síntese da Auditoria
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(8)
    p3.paragraph_format.space_after = Pt(4)
    r3 = p3.add_run("3. SÍNTESE EXECUTIVA DA AUDITORIA & NOTAS POR ÁREA FUNCIONAL")
    r3.bold = True
    r3.font.size = Pt(11.5)
    r3.font.color.rgb = RGBColor(11, 37, 69)

    doc.add_paragraph("{{sintese_narrativa}}")
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 6. Section 4: Resumo por Área Funcional (Tabela 4 Colunas)
    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(8)
    p4.paragraph_format.space_after = Pt(4)
    r4 = p4.add_run("4. RESUMO DA AUDITORIA POR ÁREA FUNCIONAL (GAPS & OPORTUNIDADES)")
    r4.bold = True
    r4.font.size = Pt(11.5)
    r4.font.color.rgb = RGBColor(11, 37, 69)

    gaps_table = doc.add_table(rows=2, cols=4)
    gaps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(gaps_table, color="CBD5E1")
    
    gap_headers = ["Área / Processo", "Gaps", "Evidência", "Descrição"]
    for c_idx, h_text in enumerate(gap_headers):
        c = gaps_table.cell(0, c_idx)
        set_cell_background(c, "0B2545")
        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    # Linha modelo com tags de repetição
    c0 = gaps_table.cell(1, 0); c0.paragraphs[0].add_run("{{gaps.area}}")
    c1 = gaps_table.cell(1, 1); c1.paragraphs[0].add_run("{{gaps.tipo_badge}}")
    c2 = gaps_table.cell(1, 2); c2.paragraphs[0].add_run("{{gaps.evidencia}}")
    c3 = gaps_table.cell(1, 3); c3.paragraphs[0].add_run("{{gaps.descricao}}")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 7. Section 5: Assinaturas
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(8)
    p5.paragraph_format.space_after = Pt(4)
    r5 = p5.add_run("5. CONSIDERAÇÕES FINAIS E ASSINATURAS")
    r5.bold = True
    r5.font.size = Pt(11.5)
    r5.font.color.rgb = RGBColor(11, 37, 69)

    doc.add_paragraph("{{conclusao_texto}}")

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_s0 = sign_table.cell(0, 0); c_s1 = sign_table.cell(0, 1)
    c_l0 = sign_table.cell(1, 0); c_l1 = sign_table.cell(1, 1)
    for c in [c_s0, c_s1, c_l0, c_l1]:
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_s0.paragraphs[0].add_run("_________________________________________\n").font.size = Pt(9)
    c_s1.paragraphs[0].add_run("_________________________________________\n").font.size = Pt(9)
    c_l0.paragraphs[0].add_run("Auditor Líder: {{auditores}}\nEquipe Auditora").bold = True
    c_l1.paragraphs[0].add_run("Representante da Direção: {{representantes}}\nUnidade Auditada").bold = True

    try:
        doc.save(template_path)
    except Exception:
        pass

    return template_path


def compute_auditoria_metricas_completas(auditoria) -> Dict[str, Any]:
    """
    Calcula com precisão matemática todos os contadores da auditoria ISO,
    utilizando a mesma lógica exata e consolidada do Dashboard de Fechamento.
    Corrige qualquer discrepância de contadores zerados.
    """
    try:
        from ..models import (
            ItemNorma,
            RespostaEntrevistaIso,
            AvaliacaoFinalRequisitoIso,
        )
    except Exception:
        ItemNorma = None
        RespostaEntrevistaIso = None
        AvaliacaoFinalRequisitoIso = None

    agendas_qs = auditoria.agendas.all() if hasattr(auditoria, 'agendas') else []
    if hasattr(agendas_qs, 'prefetch_related'):
        agendas_qs = agendas_qs.prefetch_related('perguntas', 'itens_norma', 'perguntas__itens_norma')
    agendas = list(agendas_qs)

    # Coleta de Itens do Escopo (com fallback completo)
    itens_escopo = auditoria.escopo_itens.all() if hasattr(auditoria, 'escopo_itens') else []
    if hasattr(itens_escopo, 'exists') and not itens_escopo.exists():
        if ItemNorma:
            itens_escopo = ItemNorma.objects.filter(norma=auditoria.norma)
    itens_escopo_list = list(itens_escopo)

    # Identifica itens pais que possuem subitens (avaliamos apenas folhas para não duplicar contagem)
    parent_ids = set()
    for item in itens_escopo_list:
        prefix = item.referencia + '.'
        if any(other.referencia.startswith(prefix) for other in itens_escopo_list):
            parent_ids.add(item.id)

    if RespostaEntrevistaIso:
        respostas = RespostaEntrevistaIso.objects.filter(auditoria=auditoria).prefetch_related(
            'solicitacoes', 'solicitacoes__imagens', 'pergunta__itens_norma'
        )
    else:
        respostas = []
    respostas_map = {r.pergunta_id: r for r in respostas}
    na_item_ids = set(auditoria.itens_nao_aplicaveis.values_list('id', flat=True)) if hasattr(auditoria, 'itens_nao_aplicaveis') else set()

    if AvaliacaoFinalRequisitoIso:
        avaliacoes_finais_map = {
            av.item_norma_id: av
            for av in AvaliacaoFinalRequisitoIso.objects.filter(auditoria=auditoria)
        }
    else:
        avaliacoes_finais_map = {}

    hierarchy = {"NC": 5, "P": 4, "OM": 3, "C": 2, "NA": 1}
    reverse_hierarchy = {v: k for k, v in hierarchy.items()}

    destaques_conformes = []
    pontos_a_melhorar = []
    exclusoes_na = []
    gaps_area_funcional = []

    count_c = 0
    count_obs = 0
    count_om = 0
    count_nc_menor = 0
    count_nc_maior = 0
    count_p = 0
    count_na = 0

    for item in sorted(itens_escopo_list, key=lambda x: (x.ordem or 0, natural_sort_key(x.referencia))):
        if item.id in parent_ids:
            continue

        todas_perguntas_dict = {}
        agendas_do_item = set()
        for ag in agendas:
            if item in ag.itens_norma.all():
                agendas_do_item.add(ag.titulo)
            for p in ag.perguntas.all():
                if item in p.itens_norma.all() or (not p.itens_norma.exists() and item in ag.itens_norma.all()):
                    todas_perguntas_dict[p.id] = p
                    agendas_do_item.add(ag.titulo)
                    
        # Garante que respostas/perguntas órfãs (ex: criadas via painel de revisão) sejam incluídas
        for r in respostas:
            if item in r.pergunta.itens_norma.all():
                todas_perguntas_dict[r.pergunta_id] = r.pergunta
                
        todas_perguntas_item = list(todas_perguntas_dict.values())

        agendas_do_item_objs = []
        for ag in agendas:
            if item in ag.itens_norma.all() or any(item in p.itens_norma.all() for p in ag.perguntas.all()):
                agendas_do_item_objs.append(ag)

        if not agendas_do_item_objs:
            agendas_do_item_objs = [None]

        av_final = avaliacoes_finais_map.get(item.id)

        if av_final and av_final.classificacao:
            status_item = av_final.classificacao
        elif item.id in na_item_ids:
            status_item = "NA"
        elif not todas_perguntas_item:
            status_item = "P" if any(item in ag.itens_norma.all() for ag in agendas) else "NA"
        else:
            pior_peso = 0
            for p in todas_perguntas_item:
                r = respostas_map.get(p.id)
                c = r.classificacao if r else "P"
                if c == "OBS": c = "C"
                peso = hierarchy.get(c, 2)
                if peso > pior_peso: pior_peso = peso
            status_item = reverse_hierarchy.get(pior_peso, "P")

        # Evidências globais do Requisito (usadas para destaques e contadores globais)
        evidencias_item = []
        evidencias_vistas = set()
        for p in todas_perguntas_item:
            r = respostas_map.get(p.id)
            if r:
                for s in r.solicitacoes.all():
                    if s.conclusao == 'OBS': count_obs += 1
                    ev_txt = f"{s.solicitacao.strip()}: {s.evidencia.strip()}" if (s.evidencia and s.solicitacao) else (s.evidencia.strip() if s.evidencia else s.solicitacao.strip())
                    if ev_txt and ev_txt not in evidencias_vistas:
                        evidencias_vistas.add(ev_txt)
                        evidencias_item.append(ev_txt)
                if r.texto_resposta and r.texto_resposta.strip():
                    txt = r.texto_resposta.strip()
                    if txt not in evidencias_vistas:
                        evidencias_vistas.add(txt)
                        evidencias_item.append(txt)
                        
        if av_final and av_final.justificativa and av_final.justificativa.strip():
            just_txt = f"Revisão: {av_final.justificativa.strip()}"
            if just_txt not in evidencias_vistas:
                evidencias_item.insert(0, just_txt)

        # Contadores globais (Não muda para não afetar as estatísticas)
        if status_item == 'NA':
            count_na += 1
            just_na = av_final.justificativa.strip() if av_final and av_final.justificativa else (" | ".join(evidencias_item) if evidencias_item else "Requisito Não Aplicável.")
            exclusoes_na.append({'referencia': item.referencia, 'titulo': item.titulo, 'justificativa': just_na})
        elif status_item == 'P': count_p += 1
        elif status_item == 'C': 
            count_c += 1
            destaques_conformes.append({'referencia': item.referencia, 'titulo': item.titulo, 'evidencias': evidencias_item[:3] or ["Processo conforme."]})
        elif status_item == 'OM': count_om += 1
        elif status_item == 'NC':
            is_maior = False
            if av_final and av_final.grau_nc:
                is_maior = (av_final.grau_nc == 'MAIOR')
            else:
                graus_definidos = []
                for p in todas_perguntas_item:
                    r = respostas_map.get(p.id)
                    if r:
                        if r.grau_nc: graus_definidos.append(r.grau_nc)
                        for s in r.solicitacoes.all():
                            if s.conclusao == 'NC' and s.grau_nc:
                                graus_definidos.append(s.grau_nc)
                if graus_definidos:
                    is_maior = any(g == 'MAIOR' for g in graus_definidos)
                else:
                    is_maior = any('crítica' in ev.lower() or 'grave' in ev.lower() or 'sistêmica' in ev.lower() for ev in evidencias_item)
            
            if is_maior: count_nc_maior += 1
            else: count_nc_menor += 1

        # Geração de Gaps desmembrada por Área/Agenda
        solicitacoes_processadas = set()
        for ag in agendas_do_item_objs:
            if ag:
                area_nome = ag.titulo
            else:
                ref_parts = item.referencia.split('.')
                sec_code = ref_parts[0] if ref_parts else item.referencia
                pai = ItemNorma.objects.filter(norma=auditoria.norma, referencia=sec_code).first() if ItemNorma else None
                area_nome = f"{pai.titulo if pai else item.titulo}" if not pai else f"{pai.titulo}"

            amostras_por_tipo = {'C': [], 'NC': [], 'OM': []}
            pior_peso_ag = 0
            is_maior_ag = False
            
            def add_amostra(tipo, tit, desc):
                if not tit:
                    return
                
                # Ignorar amostras genéricas se não possuírem descrição complementar
                tit_clean = tit.strip().lower()
                desc_clean = desc.strip() if desc else ""
                textos_ignorados = [
                    "avaliação realizada durante a auditoria.",
                    "avaliação realizada durante a auditoria",
                    "nova solicitação de evidência",
                    "nova solicitação de evidência.",
                    "evidências / documentos registrados"
                ]
                
                if not desc_clean and tit_clean in textos_ignorados:
                    return

                idx_desc = -1
                for i, (e_tit, e_desc) in enumerate(amostras_por_tipo[tipo]):
                    if e_desc and e_desc == desc:
                        idx_desc = i
                        break
                if idx_desc != -1:
                    e_tit, _ = amostras_por_tipo[tipo][idx_desc]
                    if tit and not e_tit:
                        amostras_por_tipo[tipo][idx_desc] = (tit, desc)
                else:
                    if (tit, desc) not in amostras_por_tipo[tipo]:
                        amostras_por_tipo[tipo].append((tit, desc))

            for p in todas_perguntas_item:
                if ag and p not in ag.perguntas.all() and item not in ag.itens_norma.all():
                    continue 
                    
                r = respostas_map.get(p.id)
                if r:
                    c_resp = r.classificacao if r.classificacao != "OBS" else "C"
                    peso_resp = hierarchy.get(c_resp, 2)
                    if peso_resp > pior_peso_ag: pior_peso_ag = peso_resp

                    for s in r.solicitacoes.all():
                        if s.id in solicitacoes_processadas:
                            continue
                            
                        if s.agenda == ag or (s.agenda is None and ag and p in ag.perguntas.all()) or (s.agenda is None and not ag):
                            solicitacoes_processadas.add(s.id)
                            c_sol = s.conclusao if s.conclusao != "OBS" else "C"
                            peso_sol = hierarchy.get(c_sol, 2)
                            if peso_sol > pior_peso_ag: pior_peso_ag = peso_sol

                            tipo = c_sol if c_sol in ['NC', 'OM'] else 'C'
                            
                            if c_sol == 'NC' and s.grau_nc == 'MAIOR':
                                is_maior_ag = True

                            tit = s.solicitacao.strip() if s.solicitacao else ""
                            desc = s.evidencia.strip() if s.evidencia else ""
                            
                            add_amostra(tipo, tit, desc)
                    
                    if (ag and p in ag.perguntas.all()) or not ag:
                        if r.texto_resposta and r.texto_resposta.strip():
                            txt = r.texto_resposta.strip()
                            tipo = c_resp if c_resp in ['NC', 'OM'] else 'C'
                            add_amostra(tipo, "", txt)

            status_ag = reverse_hierarchy.get(pior_peso_ag, "P") if pior_peso_ag > 0 else status_item
            if status_item == 'NA': status_ag = 'NA'
            if status_item == 'C' and status_ag in ['NC', 'OM']: status_ag = 'C'

            if status_ag == 'NA':
                continue
                
            desc_revisao = av_final.justificativa if (av_final and av_final.justificativa) else ""
            
            # Se a área foi avaliada mas não tem NENHUMA evidência registrada nos buckets
            has_evidence = any(amostras_por_tipo[t] for t in ['C', 'NC', 'OM'])
            if not has_evidence:
                pass # A pedido, não gera linha de amostra genérica quando não há descrição

            # Gera uma linha na tabela para cada TIPO de amostra encontrada na área
            for tipo in ['C', 'OM', 'NC']:
                amostras = amostras_por_tipo[tipo]
                
                if not amostras:
                    continue # Nenhuma amostra deste tipo
                    
                tabela_evid_parts = []
                tabela_desc_parts = []
                
                for idx, (tit, desc) in enumerate(amostras, 1):
                    prefix = f"{idx}. "
                    t_str = f"{prefix}{tit}"
                    d_str = f"{prefix}{desc}" if desc else f"{prefix}"
                    
                    tabela_evid_parts.append(t_str)
                    tabela_desc_parts.append(d_str)
                    
                separator = "\n" + "_"*30 + "\n"
                tabela_evid = separator.join(tabela_evid_parts)
                tabela_desc = separator.join(tabela_desc_parts)
                
                if desc_revisao:
                    tabela_desc += f"\n\nRevisão: {desc_revisao}"
                    
                tabela_desc = tabela_desc.strip()
                tabela_evid = tabela_evid.strip()
                
                if tipo == 'C':
                    if not tabela_desc: tabela_desc = "Evidências documentais em conformidade nesta amostra."
                    gaps_area_funcional.append({'area': area_nome, 'item_referencia': item.referencia, 'item_titulo': item.titulo, 'tipo': 'C', 'tipo_badge': "Conforme", 'descricao': desc_revisao, 'tabela_gap': 'Conforme', 'tabela_evidencia': tabela_evid, 'tabela_descricao': tabela_desc})
                elif tipo == 'OM':
                    gaps_area_funcional.append({'area': area_nome, 'item_referencia': item.referencia, 'item_titulo': item.titulo, 'tipo': 'OM', 'tipo_badge': f"Oportunidade (Item {item.referencia})", 'descricao': desc_revisao, 'tabela_gap': 'OM', 'tabela_evidencia': tabela_evid, 'tabela_descricao': tabela_desc})
                elif tipo == 'NC':
                    eh_maior = False
                    if av_final and av_final.grau_nc == 'MAIOR':
                        eh_maior = True
                    else:
                        eh_maior = is_maior_ag
                        
                    badge_nc = f"NC Maior (Item {item.referencia})" if eh_maior else f"NC Menor (Item {item.referencia})"
                    gaps_area_funcional.append({'area': area_nome, 'item_referencia': item.referencia, 'item_titulo': item.titulo, 'tipo': 'NC', 'tipo_badge': badge_nc, 'descricao': desc_revisao, 'tabela_gap': 'NC Maior' if eh_maior else 'NC Menor', 'tabela_evidencia': tabela_evid, 'tabela_descricao': tabela_desc})

    # Ordena os gaps por Área/Processo e, em seguida, pelo Item da norma (para que a Tabela de Gaps exiba agrupada)
    gaps_area_funcional.sort(key=lambda x: (x.get('area', ''), natural_sort_key(x.get('item_referencia', ''))))

    total_avaliados = count_c + count_om + count_nc_menor + count_nc_maior
    total_base_calc = total_avaliados if total_avaliados > 0 else (count_c + count_na or 1)
    percentual_conformidade = round((count_c / total_avaliados * 100), 1) if total_avaliados > 0 else (100.0 if count_c > 0 else 0.0)

    # Veredito da Auditoria
    try:
        regras = {r.status_resultado: r for r in auditoria.norma.regras_veredicto.all()}
    except Exception:
        regras = {}

    regra_apto = regras.get('APTO')
    regra_ressalva = regras.get('RESSALVA')
    regra_inapto = regras.get('INAPTO')

    gatilho_nc_maior_inapto = regra_inapto.max_nc_maior if (regra_inapto and regra_inapto.max_nc_maior is not None) else 4
    gatilho_nc_menor_inapto = regra_inapto.max_nc_menor if (regra_inapto and regra_inapto.max_nc_menor is not None) else 15
    limite_nc_maior_apto = regra_apto.max_nc_maior if (regra_apto and regra_apto.max_nc_maior is not None) else 0
    limite_nc_menor_apto = regra_apto.max_nc_menor if (regra_apto and regra_apto.max_nc_menor is not None) else 2

    # A regra deve considerar também o total de desvios (amostras individuais) gerados, 
    # e não apenas a quantidade de itens da norma afetados.
    total_amostras_nc_maior = sum(1 for g in gaps_area_funcional if g['tabela_gap'] == 'NC Maior')
    total_amostras_nc_menor = sum(1 for g in gaps_area_funcional if g['tabela_gap'] == 'NC Menor')

    if (count_nc_maior >= gatilho_nc_maior_inapto or count_nc_menor >= gatilho_nc_menor_inapto) or \
       (total_amostras_nc_maior >= gatilho_nc_maior_inapto or total_amostras_nc_menor >= gatilho_nc_menor_inapto):
        veredito_status = "INADEQUADO / NÃO CONFORME"
        veredito_cor = "DC2626"
        veredito_bg = "FEE2E2"
        veredito_parecer = regra_inapto.texto_parecer_padrao if (regra_inapto and regra_inapto.texto_parecer_padrao) else (
            "Sistema inadequado e com alto risco crítico. As evidências apontam falhas sistêmicas que inviabilizam a recomendação neste ciclo. Recomenda-se plano de ação imediato e nova auditoria."
        )
    elif (count_nc_maior <= limite_nc_maior_apto and count_nc_menor <= limite_nc_menor_apto) and \
         (total_amostras_nc_maior <= limite_nc_maior_apto and total_amostras_nc_menor <= limite_nc_menor_apto):
        veredito_status = "ADEQUADO / CONFORME"
        veredito_cor = "16A34A"
        veredito_bg = "DCFCE7"
        veredito_parecer = regra_apto.texto_parecer_padrao if (regra_apto and regra_apto.texto_parecer_padrao) else (
            "O Sistema de Gestão da Qualidade atende satisfatoriamente aos requisitos normativos e demonstra maturidade operacional para a manutenção da certificação."
        )
    else:
        veredito_status = "MELHORIA NECESSÁRIA / RESSALVA"
        veredito_cor = "D97706"
        veredito_bg = "FEF9C3"
        veredito_parecer = regra_ressalva.texto_parecer_padrao if (regra_ressalva and regra_ressalva.texto_parecer_padrao) else (
            "O Sistema de Gestão da Qualidade requer a elaboração e cumprimento de Planos de Ação Corretiva formais para os desvios pontuais identificados, sem inviabilizar a recomendação geral."
        )

    exclusoes_na.sort(key=lambda x: natural_sort_key(x['referencia']))

    return {
        'total_c': count_c,
        'pct_c': f"{(count_c/total_base_calc*100):.1f}" if total_base_calc else "0.0",
        'total_nc_menor': count_nc_menor,
        'pct_nc_menor': f"{(count_nc_menor/total_base_calc*100):.1f}" if total_base_calc else "0.0",
        'total_nc_maior': count_nc_maior,
        'pct_nc_maior': f"{(count_nc_maior/total_base_calc*100):.1f}" if total_base_calc else "0.0",
        'total_nc': count_nc_menor + count_nc_maior,
        'pct_nc': f"{((count_nc_menor + count_nc_maior)/total_base_calc*100):.1f}" if total_base_calc else "0.0",
        'total_om': count_om,
        'pct_om': f"{(count_om/total_base_calc*100):.1f}" if total_base_calc else "0.0",
        'total_obs': count_obs,
        'pct_obs': f"{(count_obs/total_base_calc*100):.1f}" if total_base_calc else "0.0",
        'total_na': count_na,
        'pct_na': "-",
        'total_p': count_p,
        'total_avaliados': total_avaliados,
        'pct_conformidade': percentual_conformidade,
        'veredito_status': veredito_status,
        'veredito_cor': veredito_cor,
        'veredito_bg': veredito_bg,
        'veredito_parecer': veredito_parecer,
        'destaques_conformes': destaques_conformes,
        'exclusoes_na': exclusoes_na,
        'gaps_area_funcional': gaps_area_funcional,
    }


def insert_paragraph_after(paragraph, text=None):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    new_p_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_p_xml)
    new_p = Paragraph(new_p_xml, paragraph._parent)
    if text:
        new_p.add_run(text)
    return new_p

def iter_all_paragraphs(doc_obj):
    for p in doc_obj.paragraphs:
        yield p
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def inject_html_to_docx(doc: Document, html_content: str, target_paragraph=None):
    """
    Converte código HTML (do editor WYSIWYG Quill/TinyMCE) nativamente em parágrafos,
    títulos, listas, tabelas e imagens no documento Word.
    """
    current_p = target_paragraph
    def add_p():
        nonlocal current_p
        if current_p is not None:
            new_p = insert_paragraph_after(current_p)
            current_p = new_p
            return new_p
        return doc.add_paragraph()

    if not html_content or not html_content.strip():
        p = add_p() if target_paragraph is None else target_paragraph
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("Nenhuma síntese ou nota registrada para este bloco.")
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)
        return

    soup = BeautifulSoup(html_content, 'html.parser')

    def process_inline(node, p):
        for child in node.children:
            if child.name is None:
                text = str(child)
                if text:
                    p.add_run(text)
            elif child.name in ['strong', 'b']:
                r = p.add_run(child.get_text())
                r.bold = True
            elif child.name in ['em', 'i']:
                r = p.add_run(child.get_text())
                r.italic = True
            elif child.name == 'u':
                r = p.add_run(child.get_text())
                r.underline = True
            elif child.name == 'a':
                r = p.add_run(child.get_text())
                r.font.color.rgb = RGBColor(37, 99, 235)
                r.underline = True
            elif child.name == 'code':
                r = p.add_run(child.get_text())
                r.font.name = 'Consolas'
                r.font.size = Pt(9.0)
            elif child.name == 'img':
                src = child.get('src', '')
                if src:
                    insert_image_to_docx(doc, src, p)
            elif child.name == 'br':
                p.add_run('\n')
            else:
                process_inline(child, p)

    def insert_image_to_docx(doc_obj, src, target_p=None):
        try:
            pil_img = None
            if src.startswith('data:image'):
                b64_data = src.split(',', 1)[1] if ',' in src else src
                img_bytes = base64.b64decode(b64_data)
                pil_img = PILImage.open(io.BytesIO(img_bytes))
            elif os.path.exists(src):
                pil_img = PILImage.open(src)

            if pil_img:
                if pil_img.mode in ('RGBA', 'P', 'LA'):
                    pil_img = pil_img.convert('RGB')
                
                max_w_px = 650
                if pil_img.width > max_w_px:
                    ratio = max_w_px / float(pil_img.width)
                    new_h = int(float(pil_img.height) * ratio)
                    pil_img = pil_img.resize((max_w_px, new_h), PILImage.Resampling.LANCZOS)

                img_io = io.BytesIO()
                pil_img.save(img_io, format='JPEG', quality=88)
                img_io.seek(0)

                img_p = target_p if target_p is not None else doc_obj.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(4)
                img_p.paragraph_format.space_after = Pt(6)
                img_p.add_run().add_picture(img_io, width=Inches(min(5.6, pil_img.width / 110.0)))
        except Exception:
            pass

    for elem in soup.find_all(recursive=False):
        if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(elem.name[1])
            p = add_p()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            
            font_size = {1: 13, 2: 12, 3: 11, 4: 10, 5: 9.5, 6: 9.5}.get(level, 11)
            r = p.add_run(elem.get_text().strip())
            r.bold = True
            r.font.size = Pt(font_size)
            r.font.color.rgb = RGBColor(15, 23, 42)

        elif elem.name in ['p', 'div']:
            p = add_p()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            process_inline(elem, p)

        elif elem.name in ['ul', 'ol']:
            is_ordered = (elem.name == 'ol')
            for idx, li in enumerate(elem.find_all('li', recursive=False), 1):
                p = add_p()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                prefix = f"{idx}. " if is_ordered else "• "
                r_prefix = p.add_run(prefix)
                r_prefix.bold = True
                r_prefix.font.color.rgb = RGBColor(37, 99, 235)
                process_inline(li, p)

        elif elem.name == 'blockquote':
            p = add_p()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(elem.get_text().strip())
            r.italic = True
            r.font.color.rgb = RGBColor(71, 85, 105)

        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if rows:
                num_cols = max(len(r.find_all(['td', 'th'])) for r in rows)
                t = doc.add_table(rows=len(rows), cols=num_cols)
                if current_p is not None:
                    current_p._p.addnext(t._tbl)
                    if t.rows:
                        current_p = t.rows[-1].cells[-1].paragraphs[-1]
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(t, color="CBD5E1")
                
                for r_idx, r_tag in enumerate(rows):
                    cells = r_tag.find_all(['td', 'th'])
                    is_header_row = (r_idx == 0 or any(c.name == 'th' for c in cells))
                    for c_idx, c_tag in enumerate(cells):
                        if c_idx < num_cols:
                            cell = t.cell(r_idx, c_idx)
                            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(1)
                            p.paragraph_format.space_after = Pt(1)
                            process_inline(c_tag, p)
                            if is_header_row:
                                set_cell_background(cell, "F1F5F9")
                                for run in p.runs:
                                    run.bold = True
                                    run.font.size = Pt(9.0)
                            else:
                                for run in p.runs:
                                    run.font.size = Pt(8.5)

        elif elem.name == 'img':
            src = elem.get('src', '')
            if src:
                insert_image_to_docx(doc, src)


def replace_tags_in_paragraph(p, tag_dict: Dict[str, str]):
    """Substitui tags simples {{tag}} preservando a formatação do parágrafo."""
    full_text = p.text
    if "{{" not in full_text:
        return

    for tag, val in tag_dict.items():
        if tag in full_text:
            full_text = full_text.replace(tag, str(val))

    p.text = full_text


def replace_tags_in_document(doc: Document, tag_dict: Dict[str, str]):
    """Varre todos os parágrafos e tabelas do documento substituindo as tags."""
    for p in doc.paragraphs:
        replace_tags_in_paragraph(p, tag_dict)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_tags_in_paragraph(p, tag_dict)


def populate_gaps_table_loop(table, gaps_list: List[Dict[str, Any]]):
    """
    Executa o Table Row Looping na tabela de Gaps (4 colunas).
    Colunas: [ Área / Processo | Gaps | Evidência | Descrição ]
    """
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    NAVY_HEX = "0B2545"
    LIGHT_BG = "F8FAFC"
    WHITE = "FFFFFF"

    hdr_row = table.rows[0]
    headers = ["Área / Processo", "Gaps", "Evidência", "Descrição"]
    for c_idx, text in enumerate(headers):
        c = hdr_row.cells[c_idx]
        set_cell_background(c, NAVY_HEX)
        set_cell_margins(c, top=80, bottom=80, left=100, right=100)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)

    if not gaps_list:
        row = table.add_row()
        for cell in row.cells:
            set_cell_margins(cell, top=70, bottom=70, left=100, right=100)
            set_cell_background(cell, WHITE)
        
        a = row.cells[0]
        b = row.cells[3]
        a.merge(b)
        p = a.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("✅ Não foram identificadas Não Conformidades ou Gaps críticos durante esta auditoria. Todos os processos auditados demonstraram atendimento aos requisitos da norma.")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(22, 101, 52)
        return

    for r_idx, gap in enumerate(gaps_list, 1):
        row = table.add_row()
        bg_color = LIGHT_BG if r_idx % 2 == 0 else WHITE

        for cell in row.cells:
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            set_cell_background(cell, bg_color)

        # 0. Área / Processo
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(f"{gap.get('area', '')}\n(Item {gap.get('item_referencia', '')})")
        r0.bold = True
        r0.font.size = Pt(8.5)
        r0.font.color.rgb = RGBColor(11, 37, 69)

        # 1. Gaps (Conforme, NC, OM, NA)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tipo_str = gap.get('tabela_gap', gap.get('tipo', ''))
        r1 = p1.add_run(tipo_str)
        r1.bold = True
        r1.font.size = Pt(8.5)
        if gap.get('tipo') == 'NC':
            r1.font.color.rgb = RGBColor(220, 38, 38)
        elif gap.get('tipo') == 'OM':
            r1.font.color.rgb = RGBColor(217, 119, 6)
        elif gap.get('tipo') == 'C':
            r1.font.color.rgb = RGBColor(22, 101, 52)
        else:
            r1.font.color.rgb = RGBColor(71, 85, 105)

        # 2. Evidência (Referência e Título do Item)
        p2 = row.cells[2].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(gap.get('tabela_evidencia', ''))
        r2.font.size = Pt(8.5)

        # 3. Descrição da avaliação (Constatação + Evidências)
        p3 = row.cells[3].paragraphs[0]
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(gap.get('tabela_descricao', ''))
        r3.font.size = Pt(8.5)



def generate_relatorio_docx_buffer(auditoria) -> io.BytesIO:
    """
    Motor principal de exportação do Relatório de Auditoria Interna DOCX.
    Utiliza Template Engine com injeção de dados no arquivo .docx cadastrado na Norma,
    com fallback automático para o template mestre da aplicação, injeção de Exclusões Justificadas,
    Síntese Narrativa WYSIWYG e Table Row Looping para os Gaps da auditoria.
    """
    # 1. Obtém o template da norma ou template mestre
    norma = auditoria.norma
    doc = None

    import base64
    if getattr(norma, 'template_docx_base64', None):
        try:
            b_data = base64.b64decode(norma.template_docx_base64)
            doc = Document(io.BytesIO(b_data))
        except Exception:
            doc = None
    elif norma.template_docx and norma.template_docx.name:
        try:
            doc = Document(norma.template_docx.path)
        except Exception:
            try:
                norma.template_docx.open('rb')
                doc = Document(norma.template_docx.file)
            except Exception:
                doc = None

    if doc is None:
        template_path = ensure_master_template_docx_exists()
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
        else:
            raise ValueError("Nenhum template encontrado. Solicite ao administrador que faça o upload na aba Modelos (Uploads).")

    # 2. Extrai métricas consolidadas e listas de dados
    dados = compute_auditoria_metricas_completas(auditoria)

    # 3. Formata variáveis de cabeçalho e metadados
    if auditoria.unidade:
        unidade_str = auditoria.unidade
    else:
        unidade_str = getattr(auditoria, 'empresa_auditada', '') or "Tecnolens Laboratório Ótico Feira Ltda"

    auditores_list = [a.get_full_name() or a.username for a in auditoria.auditores.all()]
    if auditoria.auditor_lider:
        auditor_lider_str = auditoria.auditor_lider
        if auditor_lider_str not in auditores_list:
            auditores_list.insert(0, auditor_lider_str)
    else:
        auditor_lider_str = auditores_list[0] if auditores_list else "Auditor Líder Designado"

    auditores_str = ", ".join(auditores_list) if auditores_list else (auditoria.abertura_auditores or "Equipe Auditora Designada")
    
    tipo_auditoria_val = str(getattr(auditoria, 'tipo_auditoria', 'PRESENCIAL')).upper()
    is_presencial = (tipo_auditoria_val == 'PRESENCIAL')
    is_remota = (tipo_auditoria_val == 'REMOTA')
    tipo_auditoria_str = auditoria.get_tipo_auditoria_display() if hasattr(auditoria, 'get_tipo_auditoria_display') else ("Presencial" if is_presencial else "Remota")

    dt_ini = auditoria.data_inicio.strftime('%d/%m/%Y') if auditoria.data_inicio else ""
    dt_fim = auditoria.data_fim.strftime('%d/%m/%Y') if auditoria.data_fim else ""
    datas_str = f"{dt_ini} a {dt_fim}" if (dt_ini and dt_fim and dt_ini != dt_fim) else (dt_ini or dt_fim or "Em andamento")
    rep_str = auditoria.encerramento_representantes or auditoria.abertura_representantes or "Representantes da Unidade Auditada"
    norma_desc = getattr(auditoria.norma, 'descricao', '') or ''
    escopo_str = getattr(auditoria, 'escopo', '') or "Fabricação de Lentes Oftálmicas"
    objetivo_str = getattr(auditoria, 'objetivo', '') or "Avaliar a conformidade dos processos com os requisitos da norma e a eficácia do Sistema de Gestão da Qualidade."

    conclusao_custom = getattr(auditoria, 'conclusao_texto', '') or ""
    if not conclusao_custom:
        conclusao_custom = (
            "A auditoria foi realizada com base em plano formal e amostragens representativas dos processos operacionais e de gestão. "
            "Os resultados expressos refletem a conformidade das operações em relação aos requisitos da norma e diretrizes da organização. "
            "Recomenda-se o acompanhamento dos prazos de implementação dos planos de ação para as oportunidades identificadas."
        )

    import copy
    
    # Pontos Fortes
    pontos_fortes = auditoria.pontos_fortes.all()
    pontos_fortes_linhas = []
    for pf in pontos_fortes:
        if pf.descricao:
            pontos_fortes_linhas.append(f"{pf.titulo}: {pf.descricao}")
        else:
            pontos_fortes_linhas.append(f"{pf.titulo}")
    if not pontos_fortes_linhas:
        pontos_fortes_linhas = ["Nenhum ponto forte registrado."]

    # Pontos Fracos (NC e OM)
    pontos_fracos_linhas = []
    nao_conformidades_linhas = []
    nao_conformidades_com_evidencias_linhas = []
    vistos = set()
    vistos_nc = set()
    vistos_nc_evid = set()
    for gap in dados.get('gaps_area_funcional', []):
        if gap.get('tipo') in ['NC', 'OM']:
            linha = f"{gap['tipo_badge']}: {gap['descricao']}".strip()
            # Remove trailing colon if description is empty
            if linha.endswith(':'):
                linha = linha[:-1]
            if linha not in vistos:
                vistos.add(linha)
                pontos_fracos_linhas.append(linha)
        
        if gap.get('tipo') == 'NC':
            # Formato 1: Item + Título/Descrição breve
            desc_resumida = gap.get('descricao') or gap.get('item_titulo') or ''
            linha_nc = f"{gap['tipo_badge']}: {desc_resumida}".strip()
            if linha_nc.endswith(':'):
                linha_nc = linha_nc[:-1]
            if linha_nc not in vistos_nc:
                vistos_nc.add(linha_nc)
                nao_conformidades_linhas.append(linha_nc)

            # Formato 2: Item + Evidência / Constatação detalhada (amostras)
            detalhe_evid = gap.get('tabela_descricao') or gap.get('tabela_evidencia') or gap.get('descricao') or ''
            linha_nc_detalhada = f"{gap['tipo_badge']}: {detalhe_evid}".strip()
            if linha_nc_detalhada.endswith(':'):
                linha_nc_detalhada = linha_nc_detalhada[:-1]
            if linha_nc_detalhada not in vistos_nc_evid:
                vistos_nc_evid.add(linha_nc_detalhada)
                nao_conformidades_com_evidencias_linhas.append(linha_nc_detalhada)

    if not pontos_fracos_linhas:
        pontos_fracos_linhas = ["Nenhuma Não Conformidade ou Oportunidade de Melhoria registrada."]
    if not nao_conformidades_linhas:
        nao_conformidades_linhas = ["Nenhuma Não Conformidade registrada."]
    if not nao_conformidades_com_evidencias_linhas:
        nao_conformidades_com_evidencias_linhas = ["Nenhuma Não Conformidade registrada."]

    is_adequado = (dados.get('veredito_status') == 'ADEQUADO / CONFORME')
    is_melhoria = (dados.get('veredito_status') == 'MELHORIA NECESSÁRIA / RESSALVA')
    is_inadequado = (dados.get('veredito_status') == 'INADEQUADO / NÃO CONFORME')

    # Dicionário de Injeção de Tags (com sinônimos suportados)
    tag_dict = {
        '{{unidade}}': unidade_str,
        '{{nome_unidade}}': unidade_str,
        '{{empresa_auditada}}': unidade_str,
        '{{municipio}}': "",
        '{{norma_codigo}}': auditoria.norma.codigo,
        '{{norma_descricao}}': norma_desc or "Sistema de Gestão da Qualidade",
        '{{normas_regulamentos}}': auditoria.norma.codigo,
        '{{escopo}}': escopo_str,
        '{{objetivo}}': objetivo_str,
        '{{Objetivo}}': objetivo_str,
        '{{sintese}}': getattr(auditoria, 'sintese', '') or '',
        '{{tipo_auditoria}}': tipo_auditoria_str,
        '{{modalidade}}': tipo_auditoria_str,
        '{{presencial}}': "X" if is_presencial else "   ",
        '{{remota}}': "X" if is_remota else "   ",
        '{{is_presencial}}': "X" if is_presencial else "   ",
        '{{is_remota}}': "X" if is_remota else "   ",
        '{{resultado_adequado}}': "X" if is_adequado else "   ",
        '{{resultado_melhoria}}': "X" if is_melhoria else "   ",
        '{{resultado_inadequado}}': "X" if is_inadequado else "   ",
        '{{data_inicio}}': dt_ini,
        '{{data_fim}}': dt_fim,
        '{{data_final}}': dt_fim,
        '{{data_auditoria}}': datas_str,
        '{{auditor_lider}}': auditor_lider_str,
        '{{nome_auditor_lider}}': auditor_lider_str,
        "{{municipio}}": getattr(auditoria, 'municipio', '') or "Não informado",
        '{{responsavel_qms}}': getattr(auditoria, 'responsavel_qms', '') or '',
        '{{responsável_qms}}': getattr(auditoria, 'responsavel_qms', '') or '',
        '{{auditores}}': auditores_str,
        '{{nome_auditor}}': auditores_str,
        '{{equipe_auditora}}': auditores_str,
        '{{representantes}}': rep_str,
        '{{status}}': auditoria.get_status_display() if hasattr(auditoria, 'get_status_display') else "Concluída",
        '{{data_relatorio}}': dt_fim or dt_ini or "Hoje",
        '{{total_c}}': str(dados['total_c']),
        '{{pct_c}}': str(dados['pct_c']),
        '{{total_nc_menor}}': str(dados['total_nc_menor']),
        '{{pct_nc_menor}}': str(dados['pct_nc_menor']),
        '{{total_nc_maior}}': str(dados['total_nc_maior']),
        '{{pct_nc_maior}}': str(dados['pct_nc_maior']),
        '{{total_nc}}': str(dados['total_nc']),
        '{{pct_nc}}': str(dados['pct_nc']),
        '{{total_om}}': str(dados['total_om']),
        '{{pct_om}}': str(dados['pct_om']),
        '{{total_obs}}': str(dados['total_obs']),
        '{{pct_obs}}': str(dados['pct_obs']),
        '{{total_na}}': str(dados['total_na']),
        '{{pct_na}}': str(dados['pct_na']),
        '{{total_avaliados}}': str(dados['total_avaliados']),
        '{{pct_conformidade}}': str(dados['pct_conformidade']),
        '{{taxa_conformidade}}': f"{dados['pct_conformidade']}%",
        '{{veredito_status}}': dados['veredito_status'],
        '{{veredito_parecer}}': dados['veredito_parecer'],
        '{{parecer_conclusao}}': dados['veredito_parecer'],
        '{{conclusao_parecer}}': dados['veredito_parecer'],
        '{{conclusao_texto}}': conclusao_custom,
    }

    # 4. Injeção de Tags Escalares no Documento
    replace_tags_in_document(doc, tag_dict)

    # Função auxiliar para injetar listas como parágrafos clonados mantendo a formatação (ex: bullets/numeração)
    def inject_list_tags(doc, tag, items):
        def replace_in_p(paragraph, old_text, new_text):
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return
            # Se a tag estiver dividida em múltiplos runs
            if old_text in paragraph.text:
                new_full = paragraph.text.replace(old_text, new_text)
                for i, r in enumerate(paragraph.runs):
                    if i == 0:
                        r.text = new_full
                    else:
                        r.text = ""

        for p in iter_all_paragraphs(doc):
            if tag in p.text:
                if not items:
                    items = ["Nenhum registro encontrado."]
                
                original_p_xml = copy.deepcopy(p._p)
                replace_in_p(p, tag, items[0])
                
                last_p = p
                for item_text in items[1:]:
                    new_p_xml = copy.deepcopy(original_p_xml)
                    last_p._p.addnext(new_p_xml)
                    from docx.text.paragraph import Paragraph
                    new_p = Paragraph(new_p_xml, p._parent)
                    replace_in_p(new_p, tag, item_text)
                    last_p = new_p

    # Injetar Listas formatadas nativamente
    inject_list_tags(doc, "{{pontos_fortes}}", pontos_fortes_linhas)
    inject_list_tags(doc, "{{pontos_fracos}}", pontos_fracos_linhas)
    inject_list_tags(doc, "{{pontos_fracos_nc}}", nao_conformidades_linhas)
    inject_list_tags(doc, "{{nao_conformidades}}", nao_conformidades_linhas)
    inject_list_tags(doc, "{{nao_conformidades_detalhadas}}", nao_conformidades_com_evidencias_linhas)
    inject_list_tags(doc, "{{lista_nc}}", nao_conformidades_linhas)
    inject_list_tags(doc, "{{lista_nc_detalhada}}", nao_conformidades_com_evidencias_linhas)

    # 5. Injeção da Seção "Exclusões Justificadas"
    for p in iter_all_paragraphs(doc):
        if "{{exclusoes_justificadas}}" in p.text:
            p.text = ""
            exclusoes = dados['exclusoes_na']
            if exclusoes:
                last_p = p
                for item_na in exclusoes:
                    from docx.oxml import OxmlElement
                    from docx.text.paragraph import Paragraph
                    
                    new_p_xml = OxmlElement('w:p')
                    last_p._p.addnext(new_p_xml)
                    new_p = Paragraph(new_p_xml, p._parent)
                    
                    new_p.paragraph_format.space_after = Pt(2)
                    
                    r_ref = new_p.add_run(f"• {item_na['referencia']} - ")
                    r_ref.bold = True
                    r_ref.font.size = Pt(9.5)
                    
                    r_tit = new_p.add_run(item_na['titulo'])
                    r_tit.font.size = Pt(9.5)
                    
                    last_p = new_p
            else:
                from docx.oxml import OxmlElement
                from docx.text.paragraph import Paragraph
                
                new_p_xml = OxmlElement('w:p')
                p._p.addnext(new_p_xml)
                p_na = Paragraph(new_p_xml, p._parent)
                
                p_na.paragraph_format.space_before = Pt(2)
                p_na.paragraph_format.space_after = Pt(4)
                r_na = p_na.add_run("Não foram identificadas exclusões de requisitos normativos no escopo desta auditoria. Todos os requisitos da norma foram considerados aplicáveis.")
                r_na.font.size = Pt(9.0)
                r_na.font.italic = True
                r_na.font.color.rgb = RGBColor(71, 85, 105)

    # 6. Injeção da Síntese Narrativa & Seções
    for p in iter_all_paragraphs(doc):
        if "{{sintese_narrativa}}" in p.text:
            p.text = ""
            sinteses_secoes = list(auditoria.sinteses_secao.all())
            sinteses_secoes = sorted(sinteses_secoes, key=lambda s: natural_sort_key(s.secao_referencia))
            sintese_global_html = getattr(auditoria, 'sintese', '') or ""
            tem_sintese_secao = any(bool(s.conteudo_html and s.conteudo_html.strip()) for s in sinteses_secoes)
            
            unidade_str = getattr(auditoria, 'unidade', '') or '____'
            dt_inicio = auditoria.data_inicio.strftime('%d/%m') if auditoria.data_inicio else '____'
            dt_fim = auditoria.data_fim.strftime('%d/%m/%Y') if auditoria.data_fim else '____'
            
            n_codigo = getattr(auditoria.norma, 'codigo', '')
            n_desc = getattr(auditoria.norma, 'descricao', '')
            n_str = f"{n_codigo} - {n_desc}" if n_codigo else n_desc
            
            from auditoria.models import RespostaEntrevistaIso
            from django.db.models import Q
            import re
            
            def extrair_apenas_nome(texto: str) -> str:
                if not texto: return ""
                val = str(texto).strip()
                for sep in [' - ', ' – ', ' — ', '  -  ', ' -', '- ']:
                    if sep in val:
                        val = val.split(sep, 1)[0].strip()
                for suffix in ['(Participantes)', '(Participante)', '(Entrevistados)', '(Entrevistado)', '(Auditado)', '(Auditados)']:
                    if val.endswith(suffix):
                        val = val[:-len(suffix)].strip()
                for sep in [' - ', ' – ', ' — ']:
                    if sep in val:
                        val = val.split(sep, 1)[0].strip()
                return val.strip()

            pessoas_auditadas_lista = []
            enc_reps = getattr(auditoria, 'encerramento_representantes', '')
            if enc_reps and enc_reps.strip():
                nomes_brutos = [linha.strip() for linha in re.split(r'[;\n]', enc_reps) if linha.strip()]
                nomes_vistos = set()
                for item_str in nomes_brutos:
                    item_clean = extrair_apenas_nome(item_str)
                    if not item_clean or len(item_clean) < 2: continue
                    key = item_clean.lower()
                    if key not in nomes_vistos:
                        nomes_vistos.add(key)
                        pessoas_auditadas_lista.append(item_clean)
            else:
                respostas_auditados_qs = list(
                    RespostaEntrevistaIso.objects.filter(auditoria=auditoria)
                    .filter(
                        Q(pergunta__texto_pergunta__icontains="auditadas") |
                        Q(pergunta__texto_pergunta__icontains="entrevistadas") |
                        Q(pergunta__texto_pergunta__icontains="nomes e funções") |
                        Q(pergunta__dica_auditor__icontains="participante entrevistado")
                    )
                    .select_related('pergunta')
                    .prefetch_related('solicitacoes')
                )
                nomes_vistos = set()
                for resp in respostas_auditados_qs:
                    for s in resp.solicitacoes.all():
                        ev = (s.evidencia or "").strip()
                        sol = (s.solicitacao or "").strip()
                        is_generic = sol.lower() in [
                            'entrevistado', 'entrevistados', 'pessoa auditada', 'pessoas auditadas',
                            'amostra', 'amostra #1', 'amostra #2', 'amostra #3', 'amostra #4', 'amostra #5',
                            'solicitação', 'solicitacao', ''
                        ]
                        cand = ""
                        if ev: cand = extrair_apenas_nome(ev)
                        elif sol and not is_generic: cand = extrair_apenas_nome(sol)
                        
                        if cand and len(cand) >= 2:
                            if cand.lower() not in nomes_vistos:
                                nomes_vistos.add(cand.lower())
                                pessoas_auditadas_lista.append(cand)
                    if resp.texto_resposta and resp.texto_resposta.strip():
                        for linha in re.split(r'[;\n]', resp.texto_resposta):
                            l_clean = extrair_apenas_nome(linha)
                            if l_clean and l_clean.lower() not in nomes_vistos and len(l_clean) >= 2:
                                nomes_vistos.add(l_clean.lower())
                                pessoas_auditadas_lista.append(l_clean)
            
            reps_html = "<br>".join(pessoas_auditadas_lista) if pessoas_auditadas_lista else "<i>Não preenchido</i>"
            
            intro_html = f"<p>A auditoria no {unidade_str} foi realizada no período de {dt_inicio} a {dt_fim} e durante essa auditoria foi verificado o alinhamento do QMS a norma {n_str}.</p><p><strong>Participantes:</strong><br>{reps_html}</p>"
            
            sintese_global_html = intro_html + sintese_global_html

            if sintese_global_html and sintese_global_html.strip():
                inject_html_to_docx(doc, sintese_global_html, target_paragraph=p)
                p = insert_paragraph_after(p)
                p.paragraph_format.space_after = Pt(4)

            if tem_sintese_secao:
                sec_num = 1
                for s in sinteses_secoes:
                    if s.conteudo_html and s.conteudo_html.strip():
                        p_shdr = doc.add_paragraph()
                        p_shdr.paragraph_format.space_before = Pt(8)
                        p_shdr.paragraph_format.space_after = Pt(2)
                        r_shdr = p_shdr.add_run(f"3.{sec_num} Seção {s.secao_referencia} — {s.secao_titulo}")
                        r_shdr.bold = True
                        r_shdr.font.size = Pt(10.0)
                        r_shdr.font.color.rgb = RGBColor(30, 58, 138)
                        
                        inject_html_to_docx(doc, s.conteudo_html, target_paragraph=p_shdr)
                        p = insert_paragraph_after(p_shdr)
                        p.paragraph_format.space_after = Pt(4)
                        sec_num += 1
            elif not (sintese_global_html and sintese_global_html.strip()):
                p_def = insert_paragraph_after(p)
                p_def.paragraph_format.space_after = Pt(4)
                r_def = p_def.add_run(
                    "A auditoria interna foi executada por amostragem abrangendo os processos críticos, "
                    "infraestrutura, controles operacionais e registros documentais da qualidade conforme escopo planejado."
                )
                r_def.font.size = Pt(9.0)
                r_def.font.italic = True
                r_def.font.color.rgb = RGBColor(71, 85, 105)

    # 7. Injeção na Tabela de Gaps (Table Row Looping - Tabela 4 Colunas)
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) == 4:
            first_row_text = "".join([c.text for c in table.rows[0].cells])
            if "Área" in first_row_text and "Gaps" in first_row_text:
                populate_gaps_table_loop(table, dados['gaps_area_funcional'])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
